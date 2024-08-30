import os
import docx

from abc import ABC, abstractmethod


class DocxBuilder(ABC):
    def __init__(self):
        self.path = './docx_files_storage'

    def save_doc(self, doc, name):
        try:
            if not os.path.exists(self.path):
                os.makedirs(self.path)

            path = f'{self.path}/{name}'
            doc.save(path)
            return path

        except Exception as e:
            print(f'Warning: could not create directory for docx files {e}')

    @abstractmethod
    def create(self):
        raise NotImplementedError('Method not implemented')


class ApplicationsReport(DocxBuilder):
    def create(self, data, name):
        doc = docx.Document()

        doc.add_paragraph('List:')
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'name'
        hdr_cells[1].text = 'submitted time'
        hdr_cells[2].text = 'link'

        for d in data:
            row_cells = table.add_row().cells
            row_cells[0].text = d.get('name')
            row_cells[1].text = d.get('createdTime')
            row_cells[2].text = d.get('webViewLink')

        doc.add_page_break()

        return self.save_doc(doc, name)


class ConferenceReport(DocxBuilder):
    def create(self, data, name):
        doc = docx.Document()

        doc.add_paragraph('List:')
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'name'
        hdr_cells[1].text = 'submitted time'
        hdr_cells[2].text = 'link'

        for d in data:
            row_cells = table.add_row().cells
            row_cells[0].text = d.get('name')
            row_cells[1].text = d.get('createdTime')
            row_cells[2].text = d.get('webViewLink')

        doc.add_page_break()

        return self.save_doc(doc, name)


class PublicationsReport(DocxBuilder):
    pass
